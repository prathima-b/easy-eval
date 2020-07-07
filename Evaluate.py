import csv
import spacy
from docx import Document
import os
import itertools
nlp=spacy.load('en_core_web_sm')
correct_answer=[]
student_answer=[]
correct_answers=[]
student_answers=[]
marksperq=4
document = Document("Answerkey.docx")
script=Document('Mistake.docx')
"""FOR ANSWER KEY"""
for key in document.paragraphs:
    correct_answer.append(key.text)
text=""
#print(correct_answer)
for i in correct_answer:
    text+=i
    
#print(text)
start=0
ans=""
corrected_answer=[]
for i in range(0,len(text)):
    
    if text[i]=='Q':
        
        start=0
        if ans!="":
            #print(ans)
            corrected_answer.append(ans)
            ans=""
        if text[i+1].isdigit:
            #print("qno")
            start+=1 
            continue
    elif start==1:
        start+=1
        continue
    elif start==2:
        
        ans+=text[i]
corrected_answer.append(ans)     
print("Answer Key:",corrected_answer,len(corrected_answer))   
"""FOR STUDENT ANSWER"""     
    
for key in script.paragraphs:
    student_answer.append(key.text)

text=""
#print(correct_answer)
for i in student_answer:
    text+=i
    
#print(text)
start=0
ans=""
student_answer=[]
for i in range(0,len(text)):
    
    if text[i]=='Q':
        
        start=0
        if ans!="":
            print(ans)
            student_answer.append(ans)
            ans=""
        if text[i+1].isdigit:
            #print("qno")
            start+=1 
            continue
    elif start==1:
        start+=1
        continue
    elif start==2:
        
        ans+=text[i]

student_answer.append(ans)        
#print("Student Answer",student_answer,len(student_answer))
     

    
for i in range(0,len(corrected_answer)):
    new=corrected_answer[i].replace('.','')
    new=corrected_answer[i].replace(',','')
    new=corrected_answer[i].replace('\n','')
    corrected_answer[i]=new
    
    if corrected_answer[i]!='':
        correct_answers.append(corrected_answer[i])
    
    
    new=student_answer[i].replace('.','')
    new=student_answer[i].replace(',','')
    new=student_answer[i].replace('\n','')
    student_answer[i]=new
    
    if student_answer[i]!='':
        student_answers.append(student_answer[i])

"""for (i,j) in zip(correct_answers,student_answers):
    print(i,"\n",j)
"""
f=open('AnswerKey.csv','r')
spamreader = csv.reader(f, delimiter=',', quotechar='|')
count=0
flag=0
split=0
qno=0
total=0
scores=[]
score=0
prev=0
for i in spamreader: #i is the list keywords for an answer
    
    #print(student_answer[qno])
    #Removing garbage characters
    if(flag==0):
        i[0]=i[0][3:]
        flag+=1
    length=len(i)
    
    
    #Removing empty keywords
    while(i[length-1]==''):
    
         i.pop()
         length-=1
    
    #print(i)
    #weightssss
    split=marksperq/length
    #Mark splitting for each row
    #print(split)
    
    for j in range(0,len(i)):#j is a keyword for the answer

      
        if('/' in i[j]):
            
            options=i[j].split('/')
            
            for k in options:
                
                if k in student_answers[qno]:
                    score+=split
                    print("added")
                    break
        else:
            
            
            #print("here")
            if i[j] in student_answers[qno]:
                score+=split
                print("added")
                
                
       
        
        total+=score

        score=0
    print(total)
    scores.append(total-prev)
    prev=total
    qno+=1


similarity=[]

for (i,j) in zip(correct_answers,student_answers): 
     doc1=nlp(i)
     doc2=nlp(j)
     similarity.append(doc1.similarity(doc2))
print(similarity)

for i in range(0,len(similarity)):
    if(scores[i]==0 and similarity[i]>0.970):
        total+=marksperq*0.5
        print("half added for item",i)
        scores[i]+=marksperq*0.5
    elif(similarity[i]<=0.97 and scores[i]!=0):
 
        total-=0.25*(split)
        print("reduced for item",i)
        scores[i]-=marksperq*0.25
        
print("Marks awarded:",total,"out of",16)
print(scores)
